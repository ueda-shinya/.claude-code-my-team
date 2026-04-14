"""
md-to-docx.py — Markdown → Word(.docx) 変換スクリプト
使い方: python md-to-docx.py <mdファイルパス>
出力:  同ディレクトリに <同名>.docx を生成

依存: python-docx >= 1.0.0, markdown >= 3.0
"""

import sys
import argparse
import re
import traceback
from pathlib import Path
from html.parser import HTMLParser
from html import escape

import markdown
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# ドキュメント初期設定
# ---------------------------------------------------------------------------

def setup_document() -> Document:
  """A4・余白20mm・日本語フォント設定済みのDocumentを返す"""
  doc = Document()

  # 余白設定（A4: 上下左右 20mm）
  for section in doc.sections:
    section.page_width  = Mm(210)
    section.page_height = Mm(297)
    section.top_margin    = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin   = Mm(20)
    section.right_margin  = Mm(20)

  # デフォルトスタイル: Normal
  normal_style = doc.styles['Normal']
  normal_style.font.name = 'Calibri'
  normal_style.font.size = Pt(10.5)
  # 日本語フォント設定（東アジア文字用）
  _set_east_asian_font(normal_style.font, 'Yu Gothic')

  return doc


def _set_east_asian_font(font_obj, font_name: str):
  """python-docxのFontオブジェクトに東アジア（日本語）フォントを設定する"""
  rpr = font_obj._element
  # w:rFonts 要素を取得または作成
  r_fonts = rpr.find(qn('w:rFonts'))
  if r_fonts is None:
    r_fonts = OxmlElement('w:rFonts')
    rpr.insert(0, r_fonts)
  r_fonts.set(qn('w:eastAsia'), font_name)


# ---------------------------------------------------------------------------
# テーブルのボーダー設定
# ---------------------------------------------------------------------------

def _set_table_borders(table):
  """テーブル全体にボーダーを設定する"""
  tbl = table._tbl
  tbl_pr = tbl.find(qn('w:tblPr'))
  if tbl_pr is None:
    tbl_pr = OxmlElement('w:tblPr')
    tbl.insert(0, tbl_pr)

  tbl_borders = OxmlElement('w:tblBorders')
  for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
    border_el = OxmlElement(f'w:{edge}')
    border_el.set(qn('w:val'), 'single')
    border_el.set(qn('w:sz'), '4')
    border_el.set(qn('w:color'), 'AAAAAA')
    tbl_borders.append(border_el)

  # 既存のtblBordersを置き換え
  existing = tbl_pr.find(qn('w:tblBorders'))
  if existing is not None:
    tbl_pr.remove(existing)
  tbl_pr.append(tbl_borders)


def _set_cell_shading(cell, fill_hex: str):
  """セルの背景色を設定する（fillHex例: 'E8E8E8'）"""
  tc = cell._tc
  tc_pr = tc.find(qn('w:tcPr'))
  if tc_pr is None:
    tc_pr = OxmlElement('w:tcPr')
    tc.insert(0, tc_pr)

  shd = OxmlElement('w:shd')
  shd.set(qn('w:val'), 'clear')
  shd.set(qn('w:color'), 'auto')
  shd.set(qn('w:fill'), fill_hex)

  existing = tc_pr.find(qn('w:shd'))
  if existing is not None:
    tc_pr.remove(existing)
  tc_pr.append(shd)


# ---------------------------------------------------------------------------
# インラインHTML → runへの変換
# ---------------------------------------------------------------------------

class InlineParser(HTMLParser):
  """インラインHTML（strong/em/code/a）をパースしてrun情報のリストに変換する"""

  def __init__(self):
    super().__init__()
    self._runs: list[dict] = []   # {'text': str, 'bold': bool, 'italic': bool, 'code': bool}
    self._bold   = False
    self._italic = False
    self._code   = False

  def handle_starttag(self, tag, attrs):
    if tag in ('strong', 'b'):
      self._bold = True
    elif tag in ('em', 'i'):
      self._italic = True
    elif tag == 'code':
      self._code = True

  def handle_endtag(self, tag):
    if tag in ('strong', 'b'):
      self._bold = False
    elif tag in ('em', 'i'):
      self._italic = False
    elif tag == 'code':
      self._code = False

  def handle_data(self, data):
    if data:
      self._runs.append({
        'text':   data,
        'bold':   self._bold,
        'italic': self._italic,
        'code':   self._code,
      })

  def get_runs(self) -> list[dict]:
    return self._runs


def _add_inline_runs(paragraph, html_text: str):
  """インラインHTMLをパースしてparagraphにrunを追加する"""
  # まず単純なHTMLエンティティをデコード
  parser = InlineParser()
  parser.feed(html_text)
  runs = parser.get_runs()

  if not runs:
    # フォールバック: タグなしプレーンテキスト
    plain = re.sub(r'<[^>]+>', '', html_text)
    paragraph.add_run(plain)
    return

  for run_info in runs:
    run = paragraph.add_run(run_info['text'])
    run.bold   = run_info['bold']
    run.italic = run_info['italic']
    if run_info['code']:
      run.font.name = 'Consolas'
      _set_east_asian_font(run.font, 'MS Gothic')
      run.font.size = Pt(9)


# ---------------------------------------------------------------------------
# ブロックHTML → python-docx要素への変換
# ---------------------------------------------------------------------------

class DocxBuilder(HTMLParser):
  """
  markdown.markdown()が生成したHTMLをパースして
  python-docxのDocumentに要素を順次追加するビルダー。

  情報隠蔽: 内部状態（_in_table, _rows等）は外から触れない。
  唯一の公開インターフェースは build(doc, html) のみ。
  """

  def build(self, doc: Document, html: str):
    """HTMLを受け取りdocに要素を追加する（公開インターフェース）"""
    self._doc         = doc
    self._tag_stack: list[str] = []
    self._current_para         = None
    self._current_para_tag     = None

    # テーブル状態
    self._in_table    = False
    self._rows: list[list[dict]] = []   # rows[row][col] = {'html': str, 'is_header': bool}
    self._current_row: list[dict] = []
    self._current_cell_html: str  = ''
    self._current_cell_is_header  = False

    # リスト状態
    self._list_stack: list[str] = []   # 'ul' or 'ol'
    self._list_counters: list[int] = []
    self._in_li   = False
    self._li_html = ''

    # blockquote状態
    self._in_blockquote = False
    self._bq_html = ''

    # コードブロック状態
    self._in_pre  = False
    self._in_code = False
    self._code_text = ''

    # 脚注（^[N]^ パターン）除去用
    self._current_inline = ''

    self.feed(html)
    self._flush_pending()

  # --- tag handlers ---

  def handle_data(self, data):
    if self._in_pre and self._in_code:
      self._code_text += data
    elif self._in_table:
      self._current_cell_html += data
    elif self._in_li:
      self._li_html += data
    elif self._in_blockquote:
      self._bq_html += data
    elif self._current_para_tag in ('h1', 'h2', 'h3', 'h4', 'p'):
      self._current_inline += data

  def _append_to_current(self, text: str):
    if self._in_pre and self._in_code:
      self._code_text += text
    elif self._in_table:
      self._current_cell_html += text
    elif self._in_li:
      self._li_html += text
    elif self._in_blockquote:
      self._bq_html += text
    elif self._current_para_tag in ('h1', 'h2', 'h3', 'h4', 'p'):
      self._current_inline += text

  # インラインタグ（strong/em/code/a）の開閉タグも収集対象に追加
  def handle_starttag(self, tag, attrs):
    self._tag_stack.append(tag)

    if tag in ('h1', 'h2', 'h3', 'h4'):
      self._current_para_tag = tag
      self._current_inline = ''

    elif tag == 'p':
      if not self._in_table and not self._in_li and not self._in_blockquote:
        self._current_para_tag = 'p'
        self._current_inline = ''

    elif tag in ('ul', 'ol'):
      self._list_stack.append(tag)
      self._list_counters.append(0)

    elif tag == 'li':
      self._in_li   = True
      self._li_html = ''

    elif tag == 'blockquote':
      self._in_blockquote = True
      self._bq_html = ''

    elif tag == 'pre':
      self._in_pre  = True
      self._code_text = ''

    elif tag == 'code' and self._in_pre:
      self._in_code = True

    elif tag == 'table':
      self._in_table = True
      self._rows = []

    elif tag == 'tr':
      self._current_row = []

    elif tag in ('th', 'td'):
      self._current_cell_html = ''
      self._current_cell_is_header = (tag == 'th')

    # インラインタグは収集先に追記
    elif tag in ('strong', 'b', 'em', 'i', 'code', 'a', 'sup', 'span'):
      attrs_dict = dict(attrs)
      attr_str = ''.join(f' {k}="{escape(v, quote=True)}"' for k, v in attrs_dict.items())
      self._append_to_current(f'<{tag}{attr_str}>')

  def handle_endtag(self, tag):
    if self._tag_stack and self._tag_stack[-1] == tag:
      self._tag_stack.pop()

    if tag in ('h1', 'h2', 'h3', 'h4'):
      self._flush_heading(tag, self._current_inline)
      self._current_para_tag = None
      self._current_inline = ''

    elif tag == 'p':
      if not self._in_table and not self._in_li and not self._in_blockquote:
        self._flush_paragraph(self._current_inline)
        self._current_para_tag = None
        self._current_inline = ''

    elif tag == 'li':
      self._flush_list_item()
      self._in_li   = False
      self._li_html = ''

    elif tag in ('ul', 'ol'):
      if self._list_stack:
        self._list_stack.pop()
      if self._list_counters:
        self._list_counters.pop()

    elif tag == 'blockquote':
      self._flush_blockquote(self._bq_html)
      self._in_blockquote = False
      self._bq_html = ''

    elif tag == 'pre':
      self._flush_code_block(self._code_text)
      self._in_pre  = False
      self._in_code = False
      self._code_text = ''

    elif tag == 'code' and self._in_pre:
      self._in_code = False

    elif tag == 'tr':
      if self._in_table:
        self._rows.append(self._current_row)
        self._current_row = []

    elif tag in ('th', 'td'):
      if self._in_table:
        self._current_row.append({
          'html':      self._current_cell_html,
          'is_header': self._current_cell_is_header,
        })
        self._current_cell_html = ''

    elif tag == 'table':
      self._flush_table(self._rows)
      self._in_table = False
      self._rows = []

    # インラインタグ閉じも収集先に追記
    elif tag in ('strong', 'b', 'em', 'i', 'code', 'a', 'sup', 'span'):
      self._append_to_current(f'</{tag}>')

  # --- flush helpers ---

  def _strip_inline_html(self, html_text: str) -> str:
    """脚注パターン（^[N]^等）と残余タグを除去してプレーンテキストを返す"""
    # <sup>...</sup> を除去（脚注マーカー）
    text = re.sub(r'<sup[^>]*>.*?</sup>', '', html_text, flags=re.DOTALL)
    # その他のタグを除去
    text = re.sub(r'<[^>]+>', '', text)
    return text

  def _flush_heading(self, tag: str, html_text: str):
    level = int(tag[1])  # h1→1, h2→2 ...
    text  = self._strip_inline_html(html_text).strip()
    if not text:
      return
    para = self._doc.add_heading(text, level=level)
    # 日本語フォント設定
    for run in para.runs:
      run.font.name = 'Calibri'
      _set_east_asian_font(run.font, 'Yu Gothic')

  def _flush_paragraph(self, html_text: str):
    text = html_text.strip()
    if not text:
      return
    para = self._doc.add_paragraph()
    para.style = self._doc.styles['Normal']
    _add_inline_runs(para, text)

  def _flush_list_item(self):
    html_text = self._li_html.strip()
    if not html_text:
      return

    depth = len(self._list_stack)  # 1始まり
    is_ordered = self._list_stack[-1] == 'ol' if self._list_stack else False

    if is_ordered:
      self._list_counters[-1] += 1
      style_name = 'List Number' if depth == 1 else 'List Number 2'
    else:
      style_name = 'List Bullet' if depth == 1 else 'List Bullet 2'

    try:
      para = self._doc.add_paragraph(style=style_name)
    except KeyError:
      # スタイルが存在しない場合は通常段落で代替
      para = self._doc.add_paragraph()

    _add_inline_runs(para, html_text)

  def _flush_blockquote(self, html_text: str):
    text = self._strip_inline_html(html_text).strip()
    if not text:
      return
    para = self._doc.add_paragraph()
    # インデント設定
    para.paragraph_format.left_indent = Mm(8)
    # 左罫線風: 段落の枠線（左のみ）
    _set_paragraph_left_border(para)
    run = para.add_run(text)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    run.font.italic = True

  def _flush_code_block(self, text: str):
    if not text.strip():
      return
    para = self._doc.add_paragraph()
    para.paragraph_format.left_indent = Mm(4)
    _set_paragraph_shading(para, 'F5F5F5')
    run = para.add_run(text)
    run.font.name = 'Consolas'
    _set_east_asian_font(run.font, 'MS Gothic')
    run.font.size = Pt(9)

  def _flush_table(self, rows: list):
    if not rows:
      return

    num_cols = max(len(row) for row in rows)
    num_rows = len(rows)
    if num_cols == 0:
      return

    table = self._doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    _set_table_borders(table)

    for r_idx, row in enumerate(rows):
      for c_idx, cell_info in enumerate(row):
        if c_idx >= num_cols:
          break
        cell      = table.cell(r_idx, c_idx)
        cell_html = cell_info['html'].strip()
        is_header = cell_info['is_header']

        # セル内容をクリアして書き直し（clear()は効かないためrunを明示削除）
        for r in list(cell.paragraphs[0].runs):
          r._element.getparent().remove(r._element)
        _add_inline_runs(cell.paragraphs[0], cell_html)

        # ヘッダー行: 太字 + グレー背景
        if is_header:
          _set_cell_shading(cell, 'E8E8E8')
          for para in cell.paragraphs:
            for run in para.runs:
              run.bold = True

    # テーブル幅を全幅に設定
    _set_table_width(table)

  def _flush_pending(self):
    """パース終了後に未確定の要素があれば書き出す"""
    if self._current_inline and self._current_para_tag:
      if self._current_para_tag in ('h1', 'h2', 'h3', 'h4'):
        self._flush_heading(self._current_para_tag, self._current_inline)
      else:
        self._flush_paragraph(self._current_inline)


# ---------------------------------------------------------------------------
# 段落レベルの書式設定ヘルパー
# ---------------------------------------------------------------------------

def _set_paragraph_left_border(para):
  """段落に左罫線を設定する（blockquote用）"""
  pPr = para._p.get_or_add_pPr()
  pBdr = OxmlElement('w:pBdr')
  left_el = OxmlElement('w:left')
  left_el.set(qn('w:val'),   'single')
  left_el.set(qn('w:sz'),    '12')
  left_el.set(qn('w:color'), '888888')
  left_el.set(qn('w:space'), '12')
  pBdr.append(left_el)
  pPr.append(pBdr)


def _set_paragraph_shading(para, fill_hex: str):
  """段落の背景色を設定する（コードブロック用）"""
  pPr = para._p.get_or_add_pPr()
  shd = OxmlElement('w:shd')
  shd.set(qn('w:val'),   'clear')
  shd.set(qn('w:color'), 'auto')
  shd.set(qn('w:fill'),  fill_hex)
  pPr.append(shd)


def _set_table_width(table):
  """テーブル幅を100%（全幅）に設定する"""
  tbl = table._tbl
  tbl_pr = tbl.find(qn('w:tblPr'))
  if tbl_pr is None:
    tbl_pr = OxmlElement('w:tblPr')
    tbl.insert(0, tbl_pr)

  tbl_w = OxmlElement('w:tblW')
  tbl_w.set(qn('w:w'),    '5000')
  tbl_w.set(qn('w:type'), 'pct')

  existing = tbl_pr.find(qn('w:tblW'))
  if existing is not None:
    tbl_pr.remove(existing)
  tbl_pr.append(tbl_w)


# ---------------------------------------------------------------------------
# メイン変換処理
# ---------------------------------------------------------------------------

def convert(md_path: Path) -> Path:
  """MDファイルをdocxに変換して同ディレクトリに出力する"""
  # 入力ファイル検証
  if not md_path.is_file():
    raise FileNotFoundError(f'ファイルが見つかりません: {md_path}')

  # UTF-8読み込み、失敗時CP932フォールバック
  try:
    text = md_path.read_text(encoding='utf-8')
  except UnicodeDecodeError:
    try:
      text = md_path.read_text(encoding='cp932')
    except UnicodeDecodeError as e:
      raise ValueError(
        f'ファイルの文字コードを特定できません（UTF-8 / CP932 いずれも失敗）: {md_path}'
      ) from e

  # Markdown → HTML（テーブル・フェンスドコード対応）
  md = markdown.Markdown(
    extensions=['tables', 'fenced_code'],
  )
  html_body = md.convert(text)

  # Documentを初期化
  doc = setup_document()

  # HTML → docx要素
  builder = DocxBuilder()
  builder.build(doc, html_body)

  # 出力パス（入力と同ディレクトリ固定）
  out_path = md_path.with_suffix('.docx')

  # パストラバーサル対策: 出力先が入力ファイルと同ディレクトリであることを確認
  if out_path.parent.resolve() != md_path.parent.resolve():
    raise ValueError(
      f'出力先ディレクトリが入力と異なります（パストラバーサル拒否）: {out_path.parent}'
    )

  doc.save(str(out_path))
  return out_path


# ---------------------------------------------------------------------------
# エントリポイント
# ---------------------------------------------------------------------------

def main():
  parser = argparse.ArgumentParser(
    description='Markdown → Word(.docx) 変換スクリプト'
  )
  parser.add_argument('md_file', help='変換するMarkdownファイルのパス')
  args = parser.parse_args()

  md_path = Path(args.md_file).resolve()

  if not md_path.exists():
    print(f'エラー: ファイルが見つかりません: {md_path}', file=sys.stderr)
    sys.exit(1)

  if not md_path.is_file():
    print(f'エラー: 通常ファイルではありません（ディレクトリ不可）: {md_path}', file=sys.stderr)
    sys.exit(1)

  if md_path.suffix.lower() != '.md':
    print(f'警告: .md 以外のファイルですが続行します: {md_path}', file=sys.stderr)

  print(f'変換中: {md_path}')
  try:
    out_path = convert(md_path)
  except Exception as e:
    print(f'エラー: 変換に失敗しました: {e}', file=sys.stderr)
    traceback.print_exc()
    sys.exit(1)

  print(f'出力完了: {out_path}')


if __name__ == '__main__':
  main()
